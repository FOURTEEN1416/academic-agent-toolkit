import importlib.util
import sys
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
MODULE_PATH = ROOT / "tools" / "doc_reader.py"


def load_doc_reader():
    spec = importlib.util.spec_from_file_location("doc_reader_under_test", MODULE_PATH)
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def test_main_returns_vision_failure_when_recognition_failed(tmp_path, monkeypatch):
    module = load_doc_reader()
    source = tmp_path / "input.docx"
    output = tmp_path / "report.md"
    source.write_bytes(b"placeholder")
    monkeypatch.setattr(module, "load_project_env", lambda: None)
    monkeypatch.setattr(module, "read_docx", lambda *_: {
        "text": [],
        "images": [{"index": 1, "mime": "image/png", "size": 1, "vision_error": True,
                    "content": "[视觉识别失败: unavailable]"}],
    })
    monkeypatch.setattr(sys, "argv", ["doc_reader.py", str(source), "--out", str(output)])

    assert module.main() == 3
    assert "视觉识别失败" in output.read_text(encoding="utf-8")


def test_main_allows_explicit_vision_failure_override(tmp_path, monkeypatch):
    module = load_doc_reader()
    source = tmp_path / "input.pdf"
    source.write_bytes(b"placeholder")
    monkeypatch.setattr(module, "load_project_env", lambda: None)
    monkeypatch.setattr(module, "read_pdf", lambda *_: {
        "text": [],
        "images": [{"index": 1, "mime": "image/png", "size": 1, "vision_error": True,
                    "content": "[视觉识别失败: unavailable]"}],
    })
    monkeypatch.setattr(sys, "argv", ["doc_reader.py", str(source), "--allow-vision-failure"])

    assert module.main() == 0


def test_read_docx_includes_header_images(tmp_path):
    from docx import Document
    from PIL import Image

    module = load_doc_reader()
    body_image = tmp_path / "body.png"
    header_image = tmp_path / "header.png"
    Image.new("RGB", (2, 2), "white").save(body_image)
    Image.new("RGB", (2, 2), "black").save(header_image)
    document = Document()
    document.add_picture(str(body_image))
    document.sections[0].header.paragraphs[0].add_run().add_picture(str(header_image))
    source = tmp_path / "input.docx"
    document.save(source)

    report = module.read_docx(source, use_vision=False, max_images=10)

    assert len(report["images"]) == 2
