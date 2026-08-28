#!/usr/bin/env python3
"""doc_reader.py — 完整文档读取器（防止漏读嵌入图片/照片）

背景：docx/pdf 中的关键信息（提交要求、格式规范、题目附图）经常以
嵌入图片/截图形式存在，仅提取文本会漏读。本工具读取文档时自动：
  1. 提取全部文本（段落/表格）
  2. 提取全部嵌入图片
  3. 用配置的多模态视觉模型（agnes-2.5-flash）识别每张图片内容
  4. 输出「文本 + 图片内容」合并报告，确保不遗漏

用法：
   python doc_reader.py <文件.docx|文件.pdf> [--out report.md] [--max-images N] [--no-vision] [--allow-vision-failure]
  --no-vision   不调用视觉 API（仅列出图片数量与位置，供人工查看）
  --max-images  最多识别前 N 张图片（默认全部）

输出：默认打印到 stdout；--out 时保存 markdown 报告。
退出码：0 成功；2 参数错误；3 视觉 API 不可用（--no-vision 时不受影响）。
"""
from __future__ import annotations

import argparse
import base64
import http.client
import json
import os
import ssl
import sys
from pathlib import Path
from urllib.parse import urlparse


def load_project_env() -> None:
    """加载套件 .env（视觉 API 配置）。"""
    project_root = Path(__file__).resolve().parent.parent
    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))
    try:
        from engine.env_loader import apply_env
        apply_env()
    except ImportError:
        pass


# ============ 视觉 API 调用（agnes-2.5-flash，多模态） ============

def _call_vision(image_bytes: bytes, mime: str, prompt: str) -> str:
    api_key = os.environ.get("EDITOR_AI_API_KEY") or os.environ.get("OPENAI_API_KEY")
    base_url = os.environ.get("EDITOR_AI_BASE_URL") or os.environ.get("OPENAI_BASE_URL")
    model = os.environ.get("EDITOR_AI_MODEL_ID") or os.environ.get("REVIEWER_MODEL_ID", "agnes-2.5-flash")
    if not api_key or not base_url:
        raise RuntimeError("未配置视觉 API（需要 EDITOR_AI_API_KEY / EDITOR_AI_BASE_URL）")

    b64 = base64.b64encode(image_bytes).decode("ascii")
    payload = json.dumps({
        "model": model,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
            ],
        }],
        "max_tokens": 2000,
        "stream": False,
    })

    parsed = urlparse(base_url)
    host = parsed.hostname
    path = (parsed.path or "").rstrip("/")
    if not path.endswith("/chat/completions"):
        path = path + "/chat/completions"
    scheme = parsed.scheme or "https"

    conn = (http.client.HTTPSConnection(host, 443, timeout=120, context=ssl.create_default_context())
            if scheme == "https" else http.client.HTTPConnection(host, 80, timeout=120))
    conn.request("POST", path, payload, {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    })
    resp = conn.getresponse()
    data = resp.read()
    conn.close()
    if resp.status != 200:
        raise RuntimeError(f"视觉 API HTTP {resp.status}: {data.decode('utf-8', errors='replace')[:300]}")
    result = json.loads(data.decode("utf-8"))
    return result.get("choices", [{}])[0].get("message", {}).get("content", "")


_IMAGE_PROMPT = (
    "这是一份竞赛文件中的嵌入图片（可能是截图、示意图、照片或表格图片）。"
    "请完整、逐字转录图片中的所有文字内容（标题、正文、表格、按钮、链接、界面文字等），"
    "不要遗漏任何细节；如果是示意图/照片，请描述其内容与关键信息。"
    "如果是提交要求/格式规范相关截图，请特别完整地转录所有要求条目。"
)


def _load_image_bytes(data: bytes, ext: str) -> tuple[bytes, str]:
    """返回 (字节, mime)。PDF 中的图片直接是嵌入字节。"""
    mime = {
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".gif": "image/gif", ".bmp": "image/bmp", ".tiff": "image/tiff",
    }.get(ext.lower(), "image/png")
    return data, mime


# ============ DOCX 读取 ============

def read_docx(path: Path, use_vision: bool, max_images: int) -> dict:
    from docx import Document
    doc = Document(str(path))
    report = {"text": [], "images": []}

    # 段落
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            report["text"].append(t)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                report["text"].append(" | ".join(cells))

    # 关系表包含正文 inline/floating/VML 图片；页眉和页脚有各自的 part/关系表。
    seen = set()
    parts = [doc.part]
    for section in doc.sections:
        parts.extend((section.header.part, section.footer.part))
    for part in parts:
        for rel in part.rels.values():
            target = getattr(rel, "target_part", None)
            if target is None or not hasattr(target, "blob"):
                continue
            content_type = getattr(target, "content_type", "")
            if not content_type.startswith("image/"):
                continue
            try:
                key = target.partname
                if key in seen:
                    continue
                seen.add(key)
                data = target.blob
                ext = Path(str(target.partname)).suffix
                mime = _load_image_bytes(data, ext)[1]
                info = {"index": len(report["images"]) + 1, "size": len(data), "mime": mime}
                if use_vision and len(report["images"]) < max_images:
                    try:
                        info["content"] = _call_vision(data, mime, _IMAGE_PROMPT)
                    except Exception as e:
                        info["vision_error"] = True
                        info["content"] = f"[视觉识别失败: {e}]"
                else:
                    info["content"] = "[未识别（--no-vision 或超过上限）]"
                report["images"].append(info)
            except Exception as e:
                report["images"].append({"index": len(report["images"]) + 1, "error": str(e)})
    return report


# ============ PDF 读取 ============

def read_pdf(path: Path, use_vision: bool, max_images: int) -> dict:
    import fitz
    doc = fitz.open(str(path))
    report = {"text": [], "images": []}

    for page_idx in range(len(doc)):
        page = doc[page_idx]
        text = page.get_text().strip()
        if text:
            report["text"].append(f"--- 第 {page_idx + 1} 页 ---\n{text}")

        # 页面中的嵌入图片
        for img_info in page.get_images(full=True):
            try:
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                data = base_image["image"]
                ext = base_image["ext"]
                mime = _load_image_bytes(data, "." + ext)[1]
                info = {
                    "index": len(report["images"]) + 1,
                    "page": page_idx + 1,
                    "size": len(data),
                    "mime": mime,
                }
                if use_vision and len(report["images"]) < max_images:
                    try:
                        info["content"] = _call_vision(data, mime, _IMAGE_PROMPT)
                    except Exception as e:
                        info["vision_error"] = True
                        info["content"] = f"[视觉识别失败: {e}]"
                else:
                    info["content"] = "[未识别（--no-vision 或超过上限）]"
                report["images"].append(info)
            except Exception as e:
                report["images"].append({"index": len(report["images"]) + 1, "page": page_idx + 1, "error": str(e)})

    doc.close()
    return report


# ============ 主流程 ============

def main() -> int:
    parser = argparse.ArgumentParser(description="完整文档读取器（防漏读嵌入图片）")
    parser.add_argument("file", help="docx 或 pdf 文件路径")
    parser.add_argument("--out", default="", help="输出 markdown 报告路径")
    parser.add_argument("--max-images", type=int, default=9999, help="最多识别前 N 张图片")
    parser.add_argument("--no-vision", action="store_true", help="不调用视觉 API（仅列出图片）")
    parser.add_argument("--allow-vision-failure", action="store_true",
                        help="视觉识别失败时仍返回成功（仅供人工复核场景）")
    args = parser.parse_args()

    path = Path(args.file)
    if not path.exists():
        print(f"文件不存在: {path}", file=sys.stderr)
        return 2

    load_project_env()
    use_vision = not args.no_vision
    suffix = path.suffix.lower()

    if suffix == ".docx":
        report = read_docx(path, use_vision, args.max_images)
    elif suffix == ".pdf":
        report = read_pdf(path, use_vision, args.max_images)
    else:
        print(f"不支持的文件类型: {suffix}（支持 .docx / .pdf）", file=sys.stderr)
        return 2

    # 输出
    lines = [f"# 文档读取报告: {path.name}", ""]
    lines.append(f"## 文本内容（{len(report['text'])} 段）")
    for t in report["text"]:
        lines.append(t)
        lines.append("")
    lines.append(f"## 嵌入图片（{len(report['images'])} 张）")
    if not report["images"]:
        lines.append("（无嵌入图片）")
    for img in report["images"]:
        page = f"（第{img.get('page')}页）" if img.get("page") else ""
        lines.append(f"### 图片 {img['index']}{page} [{img.get('mime', '?')} {img.get('size', 0)}B]")
        if "error" in img:
            lines.append(f"提取失败: {img['error']}")
        else:
            lines.append(img.get("content", "[无内容]"))
        lines.append("")

    output = "\n".join(lines)
    if args.out:
        Path(args.out).write_text(output, encoding="utf-8")
        print(f"报告已保存: {args.out}")
        # 同时打印摘要
        print(f"文本段数: {len(report['text'])}, 图片数: {len(report['images'])}")
    else:
        print(output)
    has_vision_error = any(image.get("vision_error") is True for image in report["images"])
    if has_vision_error and use_vision and not args.allow_vision_failure:
        print("视觉识别未完整成功；报告已标记，停止后续自动判断。", file=sys.stderr)
        return 3
    return 0


if __name__ == "__main__":
    sys.exit(main())
