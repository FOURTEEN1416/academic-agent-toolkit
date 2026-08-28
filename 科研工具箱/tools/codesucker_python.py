# -*- coding: utf-8 -*-
"""
CodeSucker Python - 软著源程序文档抽取器
Based on https://github.com/fanbuz/codesucker (Apache-2.0)
Python port of the core 5-stage pipeline: discover → clean → select → render → audit

Generates 60-page docx source code documents compliant with
China's software copyright registration requirements.
"""

import os
import re
import sys
import json
import math
import hashlib
import argparse
import datetime
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import List, Optional, Dict, Tuple
from collections import Counter

# ============================================================
# Data Types
# ============================================================

@dataclass
class CleanOptions:
    remove_comments: bool = True
    remove_blank_lines: bool = True
    mask_sensitive: bool = True
    tab_to_spaces: bool = True
    wrap_long_lines: bool = True
    wrap_columns: int = 78

@dataclass
class ProjectConfig:
    root: str = '.'
    title: str = 'My Software V1.0'
    owner: str = ''
    founded_date: str = ''
    extensions: List[str] = field(default_factory=lambda: [
        '.py', '.java', '.ts', '.js', '.go', '.rs', '.c', '.cpp', '.h',
        '.cs', '.swift', '.rb', '.php', '.lua', '.dart', '.scala', '.kt',
        '.vue', '.html', '.css', '.sql', '.sh', '.r', '.m', '.mm',
    ])
    excludes: List[str] = field(default_factory=lambda: [
        'node_modules', '.git', '__pycache__', 'venv', '.venv',
        'dist', 'build', '.idea', '.vscode', 'vendor', 'target',
    ])
    lines_per_page: int = 50
    max_pages: int = 60
    clean: CleanOptions = field(default_factory=CleanOptions)

@dataclass
class FileEntry:
    path: str
    rel_path: str
    name: str
    ext: str
    size_bytes: int
    raw_lines: int
    included: bool = True

@dataclass
class CleanedFile:
    entry: FileEntry
    lines: List[str]
    removed_comments: int = 0
    removed_blanks: int = 0
    masked_count: int = 0

@dataclass
class Selection:
    pages: List[List[str]]
    total_lines: int = 0
    picked_lines: int = 0
    truncated: bool = False

@dataclass
class AuditItem:
    check: str
    status: str  # pass, warn, fail
    detail: str = ''

@dataclass
class ProcessResult:
    cleaned: List[CleanedFile] = field(default_factory=list)
    selection: Selection = field(default_factory=Selection)
    audit_items: List[AuditItem] = field(default_factory=list)
    stats: dict = field(default_factory=dict)


# ============================================================
# Stage 1: Discover
# ============================================================

def discover(root: str, extensions: List[str], excludes: List[str]) -> List[FileEntry]:
    """Recursively scan directory for code files."""
    root_path = Path(root).resolve()
    entries = []
    exclude_dirs = set(excludes)
    
    for dirpath, dirnames, filenames in os.walk(root_path):
        # Filter excluded directories
        dirnames[:] = [d for d in dirnames if d not in exclude_dirs]
        
        for fname in filenames:
            ext = Path(fname).suffix.lower()
            if ext not in extensions:
                continue
            
            fpath = Path(dirpath) / fname
            rel = fpath.relative_to(root_path)
            
            try:
                size = fpath.stat().st_size
                # Skip binary files (>1MB or null bytes)
                if size > 2_000_000:
                    continue
                if size == 0:
                    continue
                
                # Quick binary check
                with open(fpath, 'rb') as f:
                    chunk = f.read(8192)
                    if b'\x00' in chunk:
                        continue
                
                # Count lines
                try:
                    text = fpath.read_text(encoding='utf-8', errors='replace')
                    lines = text.split('\n')
                except:
                    continue
                
                entries.append(FileEntry(
                    path=str(fpath),
                    rel_path=str(rel),
                    name=fname,
                    ext=ext,
                    size_bytes=size,
                    raw_lines=len(lines),
                ))
            except (OSError, PermissionError):
                continue
    
    # Sort by path for deterministic output
    entries.sort(key=lambda e: e.rel_path)
    return entries


# ============================================================
# Stage 2: Clean
# ============================================================

# Comment syntax for 30+ languages
COMMENT_PATTERNS = {
    '.py': ('#', '"""', "'''"),
    '.java': ('//', '/*'),
    '.kt': ('//', '/*'),
    '.ts': ('//', '/*'),
    '.js': ('//', '/*'),
    '.go': ('//', '/*'),
    '.rs': ('//', '/*'),
    '.c': ('//', '/*'),
    '.cpp': ('//', '/*'),
    '.h': ('//', '/*'),
    '.cs': ('//', '/*'),
    '.swift': ('//', '/*'),
    '.rb': ('#', '=begin'),
    '.php': ('//', '/*', '#'),
    '.lua': ('--', '--[['),
    '.dart': ('//', '/*'),
    '.scala': ('//', '/*'),
    '.vue': ('//', '/*', '#'),
    '.html': ('<!--',),
    '.css': ('/*',),
    '.sql': ('--', '/*'),
    '.sh': ('#',),
    '.r': ('#',),
    '.m': ('//', '%'),
    '.mm': ('//',),
}

# Sensitive data patterns
SENSITIVE_PATTERNS = [
    (re.compile(r'(?i)(api[_-]?key|secret[_-]?key|access[_-]?token)\s*[=:]\s*["\'][^"\']+["\']'), r'\1="REDACTED"'),
    (re.compile(r'(?i)(password|passwd|pwd)\s*[=:]\s*["\'][^"\']+["\']'), r'\1="REDACTED"'),
    (re.compile(r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b'), 'XXX.XXX.XXX.XXX'),
    (re.compile(r'1[3-9]\d{9}'), 'XXXXXXXXXXX'),
]


def clean_file(entry: FileEntry, raw_text: str, opts: CleanOptions) -> CleanedFile:
    """Clean a single file: remove comments, blanks, mask sensitive data."""
    lines = raw_text.split('\n')
    cleaned = []
    removed_comments = 0
    removed_blanks = 0
    masked_count = 0
    
    in_block_comment = False
    in_triple_quote = False
    
    for line in lines:
        original = line
        
        # Handle block comments
        if in_block_comment:
            if '*/' in line:
                line = line[line.index('*/') + 2:]
                in_block_comment = False
            else:
                removed_comments += 1
                continue
        
        if in_triple_quote:
            if '"""' in line or "'''" in line:
                marker = '"""' if '"""' in line else "'''"
                idx = line.index(marker) + 3
                line = line[idx:]
                in_triple_quote = False
            else:
                removed_comments += 1
                continue
        
        if opts.remove_comments:
            ext = entry.ext
            patterns = COMMENT_PATTERNS.get(ext, ('//', '#'))
            
            for pat in patterns:
                if pat in ('"""', "'''"):
                    count = line.count(pat)
                    if count % 2 == 1:
                        # Odd number = start/end of triple quote block
                        idx = line.rindex(pat)
                        line = line[:idx]
                        in_triple_quote = True
                        removed_comments += 1
                        break
                elif pat == '--[[':
                    if '--[[' in line:
                        idx = line.index('--[[')
                        if ']]' in line[idx:]:
                            end = line.index(']]', idx) + 2
                            line = line[:idx] + line[end:]
                        else:
                            line = line[:idx]
                            in_block_comment = True
                        removed_comments += 1
                        break
                elif pat == '=begin':
                    if '=begin' in line:
                        line = line[:line.index('=begin')]
                        in_block_comment = True
                        removed_comments += 1
                        break
                elif pat == '<!--':
                    if '<!--' in line:
                        if '-->' in line:
                            idx = line.index('<!--')
                            end = line.index('-->', idx) + 3
                            line = line[:idx] + line[end:]
                        else:
                            line = line[:line.index('<!--')]
                            in_block_comment = True
                        removed_comments += 1
                        break
                elif pat in ('//', '#', '--', '%'):
                    # Single-line comment - but not inside strings
                    idx = line.find(pat)
                    if idx >= 0:
                        # Simple heuristic: skip if inside quotes
                        before = line[:idx]
                        if before.count('"') % 2 == 0 and before.count("'") % 2 == 0:
                            line = line[:idx]
                            removed_comments += 1
        
        # Strip trailing whitespace
        line = line.rstrip()
        
        # Remove blank lines
        if opts.remove_blank_lines and not line:
            removed_blanks += 1
            continue
        
        # Tab to spaces
        if opts.tab_to_spaces:
            line = line.expandtabs(4)
        
        # Wrap long lines
        if opts.wrap_long_lines and len(line) > opts.wrap_columns:
            # CJK-aware: count CJK chars as width 2
            width = sum(2 if ord(c) > 0x2E80 else 1 for c in line)
            if width > opts.wrap_columns:
                line = line[:opts.wrap_columns]
        
        # Mask sensitive data
        if opts.mask_sensitive:
            for pattern, replacement in SENSITIVE_PATTERNS:
                new_line = pattern.sub(replacement, line)
                if new_line != line:
                    masked_count += 1
                    line = new_line
        
        cleaned.append(line)
    
    return CleanedFile(
        entry=entry,
        lines=cleaned,
        removed_comments=removed_comments,
        removed_blanks=removed_blanks,
        masked_count=masked_count,
    )


# ============================================================
# Stage 3: Select
# ============================================================

def select(files: List[CleanedFile], lines_per_page: int, max_pages: int) -> Selection:
    """Select lines for the document: full or truncated (first 1500 + last 1500)."""
    all_lines = []
    for f in files:
        all_lines.extend(f.lines)
    
    total = len(all_lines)
    max_lines = lines_per_page * max_pages
    
    if total <= max_lines:
        picked = all_lines
        truncated = False
    else:
        half = max_lines // 2
        front = all_lines[:half]
        back = all_lines[-half:]
        picked = front + back
        truncated = True
    
    # Split into pages
    pages = []
    for i in range(0, len(picked), lines_per_page):
        page = picked[i:i + lines_per_page]
        pages.append(page)
    
    return Selection(
        pages=pages,
        total_lines=total,
        picked_lines=len(picked),
        truncated=truncated,
    )


# ============================================================
# Stage 4: Render (DOCX)
# ============================================================

def render_docx(selection: Selection, config: ProjectConfig, output_path: str) -> str:
    """Render selection to DOCX format."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        # Fallback: render as TXT
        return render_txt(selection, config, output_path.replace('.docx', '.txt'))
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # Set page margins (A4)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # Add header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = config.title
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add pages
    for page_num, page in enumerate(selection.pages, 1):
        if page_num > 1:
            doc.add_page_break()
        
        for line in page:
            para = doc.add_paragraph(line)
            para.style = doc.styles['Normal']
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = Pt(14)  # Fixed line spacing
    
    doc.save(output_path)
    return output_path


def render_txt(selection: Selection, config: ProjectConfig, output_path: str) -> str:
    """Render selection to TXT format."""
    lines = []
    for page_num, page in enumerate(selection.pages, 1):
        if page_num > 1:
            lines.append('')
            lines.append(f'--- Page {page_num} ---')
            lines.append('')
        lines.extend(page)
    
    Path(output_path).write_text('\n'.join(lines), encoding='utf-8')
    return output_path


# ============================================================
# Stage 5: Audit
# ============================================================

def audit(cleaned: List[CleanedFile], selection: Selection, config: ProjectConfig) -> List[AuditItem]:
    """Run compliance checks on the output."""
    items = []
    
    # Check 1: Non-final pages have >= lines_per_page lines
    for i, page in enumerate(selection.pages[:-1], 1):
        if len(page) < config.lines_per_page:
            items.append(AuditItem(
                check=f'Page {i} line count',
                status='fail',
                detail=f'Page {i} has {len(page)} lines, expected {config.lines_per_page}',
            ))
        else:
            items.append(AuditItem(
                check=f'Page {i} line count',
                status='pass',
                detail=f'{len(page)} lines',
            ))
    
    # Check 2: Final page >= 2/3 full
    if selection.pages:
        last_page = selection.pages[-1]
        required = int(config.lines_per_page * 2 / 3)
        if len(last_page) < required:
            items.append(AuditItem(
                check='Final page fill',
                status='warn',
                detail=f'Last page has {len(last_page)} lines (< {required} = 2/3 of {config.lines_per_page})',
            ))
        else:
            items.append(AuditItem(
                check='Final page fill',
                status='pass',
                detail=f'{len(last_page)} lines',
            ))
    
    # Check 3: Total pages
    if len(selection.pages) > config.max_pages:
        items.append(AuditItem(
            check='Page count',
            status='fail',
            detail=f'{len(selection.pages)} pages exceeds max {config.max_pages}',
        ))
    else:
        items.append(AuditItem(
            check='Page count',
            status='pass',
            detail=f'{len(selection.pages)} pages',
        ))
    
    # Check 4: No residual blank lines in cleaned content
    blank_count = sum(1 for f in cleaned for l in f.lines if not l.strip())
    if blank_count > 0:
        items.append(AuditItem(
            check='Blank lines',
            status='warn',
            detail=f'{blank_count} blank lines remain after cleaning',
        ))
    else:
        items.append(AuditItem(
            check='Blank lines',
            status='pass',
            detail='No blank lines',
        ))
    
    return items


# ============================================================
# Main Pipeline
# ============================================================

def process(config: ProjectConfig, output_dir: str = '.') -> ProcessResult:
    """Run the full 5-stage pipeline."""
    print(f"[1/5] Discovering files in {config.root}...")
    entries = discover(config.root, config.extensions, config.excludes)
    print(f"  Found {len(entries)} code files")
    
    print(f"[2/5] Cleaning {len(entries)} files...")
    cleaned = []
    for entry in entries:
        try:
            text = Path(entry.path).read_text(encoding='utf-8', errors='replace')
            cf = clean_file(entry, text, config.clean)
            cleaned.append(cf)
        except Exception as e:
            print(f"  Warning: {entry.rel_path}: {e}")
    print(f"  Cleaned {len(cleaned)} files")
    
    print(f"[3/5] Selecting lines for {config.max_pages} pages...")
    sel = select(cleaned, config.lines_per_page, config.max_pages)
    print(f"  Total: {sel.total_lines} lines, Selected: {sel.picked_lines}, Pages: {len(sel.pages)}, Truncated: {sel.truncated}")
    
    print(f"[4/5] Rendering document...")
    out_path = os.path.join(output_dir, f"{config.title.replace(' ', '_')}_source_code.docx")
    try:
        render_docx(sel, config, out_path)
        print(f"  DOCX: {out_path}")
    except Exception as e:
        out_path = os.path.join(output_dir, f"{config.title.replace(' ', '_')}_source_code.txt")
        render_txt(sel, config, out_path)
        print(f"  TXT (fallback): {out_path}")
    
    print(f"[5/5] Running audit...")
    items = audit(cleaned, sel, config)
    fails = sum(1 for i in items if i.status == 'fail')
    warns = sum(1 for i in items if i.status == 'warn')
    passes = sum(1 for i in items if i.status == 'pass')
    print(f"  Audit: {passes} pass, {warns} warn, {fails} fail")
    
    for item in items:
        icon = {'pass': '✓', 'warn': '⚠', 'fail': '✗'}[item.status]
        print(f"    {icon} {item.check}: {item.detail}")
    
    # Stats
    stats = {
        'files_discovered': len(entries),
        'files_cleaned': len(cleaned),
        'total_lines': sel.total_lines,
        'picked_lines': sel.picked_lines,
        'pages': len(sel.pages),
        'truncated': sel.truncated,
        'output_file': out_path,
        'audit_pass': passes,
        'audit_warn': warns,
        'audit_fail': fails,
    }
    
    return ProcessResult(
        cleaned=cleaned,
        selection=sel,
        audit_items=items,
        stats=stats,
    )


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description='CodeSucker Python - 软著源程序文档抽取器',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python codesucker_python.py /path/to/project --title "MyApp V1.0"
  python codesucker_python.py /path/to/project --title "MyApp V1.0" --owner "My Company"
  python codesucker_python.py /path/to/project --extensions .py .java --max-pages 30
        """
    )
    parser.add_argument('root', help='项目根目录')
    parser.add_argument('--title', default='My Software V1.0', help='软件全称+版本号')
    parser.add_argument('--owner', default='', help='著作权人（用于署名校验）')
    parser.add_argument('--output', default='.', help='输出目录')
    parser.add_argument('--extensions', nargs='+', default=None, help='文件扩展名')
    parser.add_argument('--excludes', nargs='+', default=None, help='排除目录')
    parser.add_argument('--max-pages', type=int, default=60, help='最大页数')
    parser.add_argument('--lines-per-page', type=int, default=50, help='每页行数')
    parser.add_argument('--no-clean', action='store_true', help='不清理注释')
    parser.add_argument('--json', action='store_true', help='JSON输出')
    
    args = parser.parse_args()
    
    config = ProjectConfig(
        root=args.root,
        title=args.title,
        owner=args.owner,
        max_pages=args.max_pages,
        lines_per_page=args.lines_per_page,
    )
    
    if args.extensions:
        config.extensions = args.extensions
    if args.excludes:
        config.excludes = args.excludes
    if args.no_clean:
        config.clean.remove_comments = False
        config.clean.remove_blank_lines = False
    
    result = process(config, args.output)
    
    if args.json:
        print(json.dumps(result.stats, indent=2, ensure_ascii=False))
    
    return 0 if all(i.status != 'fail' for i in result.audit_items) else 1


if __name__ == '__main__':
    sys.exit(main())
